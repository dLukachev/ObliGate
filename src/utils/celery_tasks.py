import asyncio
import re
import os
import logging
import pymupdf
from docx import Document
from src.utils.celery_client import app
from src.utils.redis_client import get_redis_session
from src.data.db.base import get_db_session
from src.repositories.contract_repo import ContractRepository

from src.docs_checker.check_file import get_and_send_to_llm
from src.utils.bot_for_remind import send_remind_in_telegram

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s: %(levelname)s/%(processName)s: %(message)s')
logger = logging.getLogger(__name__)

@app.task
async def async_set_redis(key: str, value: str, expire: int = 10):
    async with get_redis_session() as redis:
        await redis.set(key, value, ex=expire)
        return await redis.get(key)

@app.task
def remind(message):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    loop.run_until_complete(send_remind_in_telegram(message))
    loop.close()

@app.task
def file_extract(document: int):
    result = get_and_send_to_llm(document)
    return result

@app.task
def process_document(document_id: int, file_path: str):
    db = next(get_db_session())
    repo = ContractRepository(db)

    if not os.path.exists(file_path):
        logger.error(f"Файл {file_path} не найден")
        return None

    file_extension = os.path.splitext(file_path)[1].lower()
    logger.info(f"Обработка файла с расширением: {file_extension}")
    full_text = ""

    if file_extension == ".pdf":
        try:
            doc = pymupdf.open(file_path)
            logger.info(f"Открыт PDF: {file_path}, страниц: {len(doc)}")
            
            paragraph_index = 0
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text("text") # type: ignore
                full_text += page_text
                logger.info(f"Страница {page_num + 1}: извлечено {len(page_text)} символов")
                
                blocks = page.get_text("dict")["blocks"] # type: ignore
                for block in blocks:
                    if block["type"] == 0:
                        run_index = 0
                        for line in block["lines"]:
                            for span in line["spans"]:
                                span_text = span["text"].strip()
                                if span_text:
                                    bbox = span["bbox"]
                                    logger.info(f"Span: текст='{span_text}', bbox={bbox}, page={page_num + 1}")
                                    
                                    # Фильтрация: сохраняем, если есть хотя бы одна буква или цифра
                                    if re.search(r'[a-zA-Zа-яА-ЯЁё0-9]', span_text):
                                        repo.create_citation(
                                            document_id=document_id,
                                            text=span_text,
                                            page=page_num + 1,
                                            bbox=bbox,
                                            paragraph_index=paragraph_index,
                                            run_index=run_index
                                        )
                                    else:
                                        logger.info(f"Пропущен span: '{span_text}' — нет букв или цифр")
                                    run_index += 1
                        paragraph_index += 1
                
                if not page_text:
                    dict_text = page.get_text("dict") # type: ignore
                    logger.info(f"Страница {page_num + 1} (dict): {dict_text.get('blocks', [])}")
            
            doc.close()
            logger.info(f"Закрыт документ. Извлеченный текст длиной: {len(full_text)} символов")
            return full_text

        except Exception as e:
            logger.error(f"Ошибка при чтении PDF {file_path}: {type(e).__name__}: {str(e)}")
            return None

    elif file_extension == ".docx":
        try:
            doc = Document(file_path)
            full_text = ""
            paragraph_index = 0

            # Обработка параграфов
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    full_text += para_text + "\n"
                    logging.info(f"Параграф {paragraph_index}: текст='{para_text}'")
                    
                    # Обработка runs внутри параграфа для более детальной трассировки
                    run_index = 0
                    for run in para.runs:
                        run_text = run.text.strip()
                        if run_text:
                            logging.info(f"Run {run_index} в параграфе {paragraph_index}: текст='{run_text}'")
                            # Фильтрация: сохраняем, если есть хотя бы одна буква или цифра
                            if re.search(r'[a-zA-Zа-яА-ЯЁё0-9]', run_text):
                                repo.create_citation(
                                    document_id=document_id,
                                    text=run_text,
                                    page=0,
                                    bbox=None,
                                    paragraph_index=paragraph_index,
                                    run_index=run_index
                                )
                            else:
                                logger.info(f"Пропущен run: '{run_text}' — нет букв или цифр")
                            run_index += 1
                    paragraph_index += 1

            # Обработка таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            full_text += cell_text + "\n"
                            logging.info(f"Таблица, ячейка: текст='{cell_text}', paragraph_index={paragraph_index}")
                            # Сохраняем как параграф, но обрабатываем runs внутри ячейки
                            run_index = 0
                            for run in cell.paragraphs[0].runs:  # Предполагаем, что в ячейке один параграф
                                run_text = run.text.strip()
                                if run_text:
                                    logging.info(f"Run {run_index} в ячейке: текст='{run_text}'")
                                    # Фильтрация: сохраняем, если есть хотя бы одна буква или цифра
                                    if re.search(r'[a-zA-Zа-яА-ЯЁё0-9]', run_text):
                                        repo.create_citation(
                                            document_id=document_id,
                                            text=run_text,
                                            page=0,
                                            bbox=None,
                                            paragraph_index=paragraph_index,
                                            run_index=run_index
                                        )
                                    else:
                                        logger.info(f"Пропущен run: '{run_text}' — нет букв или цифр")
                                    run_index += 1
                        paragraph_index += 1

            logging.info(f"Извлечен текст из DOCX длиной: {len(full_text)} символов")

            # ставим таску на обработку документа
            file_extract.delay(document_id) # type: ignore

            return full_text

        except Exception as e:
            logging.error(f"Ошибка при чтении DOCX {file_path}: {type(e).__name__}: {str(e)}")
            return None

    else:
        logging.error(f"Неподдерживаемое расширение файла: {file_extension}")
        return None